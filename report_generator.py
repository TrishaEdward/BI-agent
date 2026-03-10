from docx import Document
from query_agent import analyze_business

TEMPLATE_PATH = "BI report.docx"


def generate_report(question, selected_sections):

    # Get AI insight + business data
    ai_answer, insights = analyze_business(question)

    # Load template
    doc = Document(TEMPLATE_PATH)

    # Report Title
    title = doc.add_heading("Business Intelligence Report", level=1)
    title.alignment = 1   # Center align title

    # -----------------------------
    # Executive Summary
    # -----------------------------
    if "Executive Summary" in selected_sections:

        doc.add_heading("Executive Summary", level=2)

        p = doc.add_paragraph()
        p.add_run(ai_answer)


    # -----------------------------
    # Key Insights
    # -----------------------------
    if "Key Insights" in selected_sections:

        doc.add_heading("Key Insights", level=2)

        for sector, value in insights["pipeline_by_sector"].items():

            doc.add_paragraph(
                f"{sector} sector pipeline value: {value:,.0f}",
                style="List Bullet"
            )


    # -----------------------------
    # Pipeline Overview
    # -----------------------------
    if "Pipeline Overview" in selected_sections:

        doc.add_heading("Pipeline Overview", level=2)

        doc.add_paragraph(
            f"Total pipeline value: {insights['total_pipeline']:,.0f}"
        )


    # -----------------------------
    # Sector Performance
    # -----------------------------
    if "Sector Performance" in selected_sections:

        doc.add_heading("Sector Performance", level=2)

        for sector, value in insights["revenue_by_sector"].items():

            doc.add_paragraph(
                f"{sector} generated revenue of {value:,.0f}",
                style="List Bullet"
            )


    # -----------------------------
    # Deal Stage Analysis
    # -----------------------------
    if "Deal Stage Analysis" in selected_sections:

        doc.add_heading("Deal Stage Analysis", level=2)

        for stage, count in insights["deals_by_stage"].items():

            doc.add_paragraph(
                f"{stage}: {count} deals",
                style="List Bullet"
            )


    # -----------------------------
    # Visualisations
    # -----------------------------
    if "Visualisations" in selected_sections:

        doc.add_heading("Visualisations", level=2)

        doc.add_paragraph(
            "Refer to the dashboard for visual charts showing pipeline distribution, "
            "sector performance, and deal stage trends."
        )


    # -----------------------------
    # Statistical Data
    # -----------------------------
    if "Statistical Data" in selected_sections:

        doc.add_heading("Statistical Data", level=2)

        doc.add_paragraph(
            f"Total Pipeline Value: {insights['total_pipeline']:,.0f}"
        )

        for prob, count in insights["deals_by_probability"].items():

            doc.add_paragraph(
                f"{prob} probability deals: {count}",
                style="List Bullet"
            )


    # -----------------------------
    # Recommendations
    # -----------------------------
    if "Recommendations" in selected_sections:

        doc.add_heading("Recommendations", level=2)

        doc.add_paragraph(
            "Leadership should focus on converting proposal-stage deals into negotiations "
            "and prioritizing high probability opportunities."
        )

    # Save report
    file_name = "Generated_BI_Report.docx"

    doc.save(file_name)

    return file_name


# =====================================================
# TEST SECTION
# =====================================================

if __name__ == "__main__":

    test_question = "Provide an overview of the current business pipeline."

    test_sections = [
        "Executive Summary",
        "Key Insights",
        "Pipeline Overview",
        "Sector Performance",
        "Deal Stage Analysis",
        "Visualisations",
        "Statistical Data",
        "Recommendations"
    ]

    report_file = generate_report(test_question, test_sections)

    print("Test complete. Report generated:", report_file)